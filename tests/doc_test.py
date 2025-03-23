import pytest
from docx import Document
from parser.doc import parse
import os

@pytest.fixture
def temp_doc():
    doc = Document()
    doc.add_paragraph('Test paragraph 1')
    doc.add_paragraph('')  # empty paragraph
    doc.add_paragraph('Test paragraph 2')
    test_file = 'test_document.docx'
    doc.save(test_file)
    yield test_file
    # Cleanup
    os.remove(test_file)

def test_parse_document(temp_doc):
    result = parse(temp_doc)
    assert result == 'Test paragraph 1\nTest paragraph 2'

def test_parse_empty_document():
    doc = Document()
    test_file = 'empty_document.docx'
    doc.save(test_file)
    try:
        result = parse(test_file)
        assert result == ''
    finally:
        os.remove(test_file)

def test_parse_multiple_paragraphs(temp_doc):
    result = parse(temp_doc)
    assert len(result.split('\n')) == 2
    assert 'Test paragraph 1' in result
    assert 'Test paragraph 2' in result